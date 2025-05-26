import cv2
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os
import tempfile

key_frame_counter_list = [2, 4, 6, 8, 10, 20]
dzielniki = [1, 2, 4, 8]
subsamplings = ["4:4:4", "4:2:2", "4:4:0", "4:2:0", "4:1:1", "4:1:0"]
ile = 10

pliki = [
    ("clip_1.mp4", False),
    ("clip_1.mp4", True)
]

doc = Document()
doc.add_heading("Wyniki kompresji wideo", 0)

temp_dir = tempfile.mkdtemp()

class data:
    def __init__(self):
        self.Y = None
        self.Cb = None
        self.Cr = None

def Chroma_subsampling(L, subsampling):
    if subsampling == "4:4:4": return L
    elif subsampling == "4:2:2": return L[:, ::2]
    elif subsampling == "4:4:0": return L[::2, :]
    elif subsampling == "4:2:0": return L[::2, ::2]
    elif subsampling == "4:1:1": return L[:, ::4]
    elif subsampling == "4:1:0": return L[::2, ::4]
    return L

def Chroma_resampling(L, subsampling):
    if L is None: raise ValueError("Chroma_resampling: received None matrix")
    if subsampling == "4:4:4": return L
    elif subsampling == "4:2:2": return np.repeat(L, 2, axis=1)
    elif subsampling == "4:4:0": return np.repeat(L, 2, axis=0)
    elif subsampling == "4:2:0": return np.repeat(np.repeat(L, 2, axis=0), 2, axis=1)
    elif subsampling == "4:1:1": return np.repeat(L, 4, axis=1)
    elif subsampling == "4:1:0": return np.repeat(np.repeat(L, 2, axis=0), 4, axis=1)
    return L

def frame_image_to_class(frame, subsampling):
    f = data()
    f.Y = frame[:, :, 0].astype(int)
    f.Cb = Chroma_subsampling(frame[:, :, 2].astype(int), subsampling)
    f.Cr = Chroma_subsampling(frame[:, :, 1].astype(int), subsampling)
    return f

def frame_layers_to_image(Y, Cr, Cb, subsampling):
    Cb = Chroma_resampling(Cb, subsampling)
    Cr = Chroma_resampling(Cr, subsampling)
    return np.dstack([Y, Cr, Cb]).clip(0, 255).astype(np.uint8)

def compress_KeyFrame(f):
    k = data()
    k.Y = f.Y
    k.Cb = f.Cb
    k.Cr = f.Cr
    return k

def decompress_KeyFrame(k, subsampling):
    return frame_layers_to_image(k.Y, k.Cr, k.Cb, subsampling)

def compress_not_KeyFrame(f, k, dzielnik):
    c = data()
    c.Y = (f.Y - k.Y) // dzielnik
    c.Cb = (f.Cb - k.Cb) // dzielnik
    c.Cr = (f.Cr - k.Cr) // dzielnik
    return c

def decompress_not_KeyFrame(c, k, dzielnik, subsampling):
    Y = c.Y * dzielnik + k.Y
    Cb = c.Cb * dzielnik + k.Cb
    Cr = c.Cr * dzielnik + k.Cr
    return frame_layers_to_image(Y, Cr, Cb, subsampling)

def save_difference_image(original, decompressed, title):
    diff = cv2.absdiff(original, decompressed)
    rgb_original = cv2.cvtColor(original, cv2.COLOR_YCrCb2RGB)
    rgb_decompressed = cv2.cvtColor(decompressed, cv2.COLOR_YCrCb2RGB)
    rgb_diff = cv2.cvtColor(diff, cv2.COLOR_YCrCb2RGB)
    fig, axs = plt.subplots(1, 3, figsize=(9, 3))
    axs[0].imshow(rgb_original)
    axs[0].set_title("Oryginał")
    axs[1].imshow(rgb_decompressed)
    axs[1].set_title("Dekompresja")
    axs[2].imshow(rgb_diff)
    axs[2].set_title("Różnica")
    for ax in axs: ax.axis("off")
    plt.tight_layout()
    path = os.path.join(temp_dir, title.replace(" ", "_") + ".png")
    plt.savefig(path)
    plt.close(fig)
    return path

for subsampling in subsamplings:
    for dzielnik in dzielniki:
        for key_frame_counter in key_frame_counter_list:
            for plik, use_rle in pliki:
                cap = cv2.VideoCapture(plik)
                ret, frame = cap.read()
                if not ret:
                    continue
                frame = cv2.cvtColor(frame, cv2.COLOR_BGR2YCrCb)
                first = frame_image_to_class(frame, subsampling)
                key = compress_KeyFrame(first)
                decompressed = decompress_KeyFrame(key, subsampling)
                title = f"KeyFrame {plik} {subsampling} dz={dzielnik}"
                img_path = save_difference_image(frame, decompressed, title)
                doc.add_paragraph(title)
                doc.add_picture(img_path, width=Inches(6))
                for i in range(1, ile):
                    ret, frame = cap.read()
                    if not ret:
                        break
                    frame = cv2.cvtColor(frame, cv2.COLOR_BGR2YCrCb)
                    f = frame_image_to_class(frame, subsampling)
                    if i % key_frame_counter == 0:
                        key = compress_KeyFrame(f)
                        decompressed = decompress_KeyFrame(key, subsampling)
                    else:
                        c = compress_not_KeyFrame(f, key, dzielnik)
                        decompressed = decompress_not_KeyFrame(c, key, dzielnik, subsampling)
                    title = f"{plik} frame {i} {subsampling} dz={dzielnik}"
                    img_path = save_difference_image(frame, decompressed, title)
                    doc.add_paragraph(title)
                    doc.add_picture(img_path, width=Inches(6))
                cap.release()

doc.save("video_compression_results.docx")