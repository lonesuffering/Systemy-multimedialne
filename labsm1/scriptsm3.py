import soundfile as sf
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from io import BytesIO
import scipy.fftpack

def analyze_audio(file, fsize):
    data, fs = sf.read(file)
    max_freq = None
    max_amp = 0
    yf = scipy.fftpack.fft(data, fsize)
    freq_axis = np.linspace(0, fs / 2, fsize // 2)
    magnitude = np.abs(yf[:fsize // 2])
    max_idx = np.argmax(magnitude)
    max_freq = freq_axis[max_idx]
    max_amp = magnitude[max_idx]
    return max_freq, max_amp

def generate_report():
    document = Document()
    document.add_heading('Audio Analysis Report', 0)

    files = ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav']
    fsizes = [2 ** 8, 2 ** 12, 2 ** 16]

    for file in files:
        document.add_heading(f'Analysis of {file}', level=1)

        for fsize in fsizes:
            fig, axs = plt.subplots(2, 1, figsize=(10, 6))
            data, fs = sf.read(file)
            axs[0].plot(data)
            axs[0].set_title(f'Time Domain (FFT size={fsize})')
            axs[0].set_xlabel('Sample')
            axs[0].set_ylabel('Amplitude')
            yf = scipy.fftpack.fft(data, fsize)
            freq_axis = np.linspace(0, fs / 2, fsize // 2)
            epsilon = 1e-10
            magnitude_db = 20 * np.log10(np.abs(yf[:fsize // 2]) + epsilon)
            axs[1].plot(freq_axis, magnitude_db)
            axs[1].set_title('Frequency Domain')
            axs[1].set_xlabel('Frequency [Hz]')
            axs[1].set_ylabel('Magnitude [dB]')

            plt.tight_layout()

            memfile = BytesIO()
            fig.savefig(memfile)
            plt.close(fig)

            document.add_heading(f'FFT Size: {fsize}', level=2)
            document.add_picture(memfile, width=Inches(6))

            max_freq, max_amp = analyze_audio(file, fsize)
            document.add_paragraph(
                f'Highest peak at: {max_freq:.2f} Hz\n'
                f'Amplitude: {max_amp:.2f} dB'
            )

    document.save('audio_analysis.docx')
    print("Report generated: audio_analysis.docx")


if __name__ == "__main__":
    generate_report()