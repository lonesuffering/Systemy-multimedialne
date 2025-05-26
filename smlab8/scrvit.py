import numpy as np
import matplotlib.pyplot as plt
import cv2
from tqdm import tqdm
import scipy.fftpack
import os
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches

# rle
def count_repetitions(data, start_idx):
    if start_idx >= len(data):
        return 0
        
    symbol = data[start_idx]
    count = 1
    i = start_idx + 1
    
    while i < len(data) and data[i] == symbol:
        count += 1
        i += 1
    
    return count

def count_differences(data, start_idx):
    if start_idx >= len(data):
        return 0
        
    count = 1
    i = start_idx + 1
    
    while i < len(data) - 1 and data[i] != data[i+1]:
        count += 1
        i += 1
    
    if i == len(data) - 1:
        count += 1
    
    return count

def compress_rle(data):
    dims = np.array([len(data.shape)], dtype=int)
    shape_info = np.concatenate([dims, data.shape])
    
    flattened = data.flatten()
    
    buffer = np.zeros(flattened.size * 2, dtype=int)
    buffer_idx = 0
    
    i = 0
    with tqdm(total=len(flattened), desc="RLE Compression") as pbar:
        while i < len(flattened):
            repetitions = count_repetitions(flattened, i)
            buffer[buffer_idx] = repetitions
            buffer_idx += 1
            buffer[buffer_idx] = flattened[i]
            buffer_idx += 1
            i += repetitions
            pbar.update(repetitions)
    
    compressed_data = buffer[:buffer_idx]
    
    
    compressed = np.concatenate([shape_info, compressed_data])
    
    return compressed

def decompress_rle(compressed):
    dims_count = compressed[0]
    shape = tuple(compressed[1:dims_count+1])
    compressed_data = compressed[dims_count+1:]
    
    decompressed_size = np.prod(shape)
    decompressed = np.zeros(decompressed_size, dtype=int)
    
    decompressed_idx = 0
    i = 0
    with tqdm(total=decompressed_size, desc="RLE Decompression") as pbar:
        while i < len(compressed_data):
            repetitions = compressed_data[i]
            symbol = compressed_data[i+1]
            decompressed[decompressed_idx:decompressed_idx+repetitions] = symbol
            decompressed_idx += repetitions
            i += 2
            pbar.update(repetitions)
    
    
    return decompressed.reshape(shape)





class ver2:
    def __init__(self,Y,Cb,Cr,OGShape,Ratio="4:4:4",QY=np.ones((8,8)),QC=np.ones((8,8)), stats=None):
        self.shape = OGShape
        self.Y=Y
        self.Cb=Cb
        self.Cr=Cr
        self.ChromaRatio=Ratio
        self.QY=QY
        self.QC=QC
        self.stats = stats
        
def RGBtoYCrCb(input):
    return cv2.cvtColor(input,cv2.COLOR_RGB2YCrCb).astype(int)

def YCrCbtoRGB(input):
    return cv2.cvtColor(input.astype(np.uint8),cv2.COLOR_YCrCb2RGB)

def dct2(a):
    return scipy.fftpack.dct(scipy.fftpack.dct(a.astype(float) - 128, axis=0, norm='ortho'), axis=1, norm='ortho')

def idct2(a):
    return scipy.fftpack.idct(scipy.fftpack.idct(a, axis=0, norm='ortho'), axis=1, norm='ortho') + 128

def zigzag(A):
    template= np.array([
            [0,  1,  5,  6,  14, 15, 27, 28],
            [2,  4,  7,  13, 16, 26, 29, 42],
            [3,  8,  12, 17, 25, 30, 41, 43],
            [9,  11, 18, 24, 31, 40, 44, 53],
            [10, 19, 23, 32, 39, 45, 52, 54],
            [20, 22, 33, 38, 46, 51, 55, 60],
            [21, 34, 37, 47, 50, 56, 59, 61],
            [35, 36, 48, 49, 57, 58, 62, 63],
            ])
    if len(A.shape)==1:
        B=np.zeros((8,8))
        for r in range(0,8):
            for c in range(0,8):
                B[r,c]=A[template[r,c]]
    else:
        B=np.zeros((64,))
        for r in range(0,8):
            for c in range(0,8):
                B[template[r,c]]=A[r,c]
    return B

def CompressBlock(block, Q):
    d = dct2(block)
    qd = np.round(d / Q).astype(int)

    return zigzag(qd)

def DecompressBlock(vector, Q):
    qd = zigzag(vector)
    d = qd * Q
    return idct2(d)


def CompressLayer(L, Q):
    S = np.array([])
    for w in range(0, L.shape[0], 8):
        for k in range(0, L.shape[1], 8):
            block = L[w:(w+8), k:(k+8)]
            S = np.append(S, CompressBlock(block, Q))
    
    return S


def DecompressLayer(S, shape, Q):
    L = np.zeros(shape, dtype=float)
    
    block_count_w = shape[0] // 8
    block_count_h = shape[1] // 8
    
    for idx in range(block_count_w * block_count_h):
        vector = S[idx*64:(idx+1)*64]
        m = shape[1] // 8
        k = int((idx % m) * 8)
        w = int((idx // m) * 8)
        L[w:(w+8), k:(k+8)] = DecompressBlock(vector, Q)
        
    return L


def chromaSubsample(channel, ratio):
    return channel[:, ::2] if ratio == "4:2:2" else channel

def chromaUpsample(channel, ratio):
   return np.repeat(channel, 2, axis=1) if ratio == "4:2:2" else channel

def compressJPEG(RGB, Ratio="4:4:4", QY=np.ones((8,8)), QC=np.ones((8,8))):
    YCrCb = RGBtoYCrCb(RGB)
    Y, Cr, Cb = YCrCb[:,:,0], YCrCb[:,:,1], YCrCb[:,:,2]

    Cr_sub = chromaSubsample(Cr, Ratio)
    Cb_sub = chromaSubsample(Cb, Ratio)

    Y_dct = CompressLayer(Y, QY)
    Cr_dct = CompressLayer(Cr_sub, QC)
    Cb_dct = CompressLayer(Cb_sub, QC)

    Y_c = compress_rle(Y_dct)
    Cr_c = compress_rle(Cr_dct)
    Cb_c = compress_rle(Cb_dct)


    stats = {}
    
    y_original_size = Y_dct.size 
    y_compressed_size = Y_c.size
    y_ratio = 100 * y_compressed_size / y_original_size
    stats['Y'] = {
        'original_size': y_original_size,
        'compressed_size': y_compressed_size,
        'compression_ratio': y_ratio
    }
    
    cr_original_size = Cr_dct.size 
    cr_compressed_size = Cr_c.size
    cr_ratio = 100 * cr_compressed_size / cr_original_size
    stats['Cr'] = {
        'original_size': cr_original_size,
        'compressed_size': cr_compressed_size,
        'compression_ratio': cr_ratio
    }

    cb_original_size = Cb_dct.size 
    cb_compressed_size = Cb_c.size
    cb_ratio = 100 * cb_compressed_size / cb_original_size
    stats['Cb'] = {
        'original_size': cb_original_size,
        'compressed_size': cb_compressed_size,
        'compression_ratio': cb_ratio
    }

    total_original = y_original_size + cr_original_size + cb_original_size
    total_compressed = y_compressed_size + cr_compressed_size + cb_compressed_size
    total_ratio = 100 * total_compressed / total_original
    stats['Total'] = {
        'original_size': total_original,
        'compressed_size': total_compressed,
        'compression_ratio': total_ratio
    }
    
    y_original_size = Y.size 
    y_compressed_size = compress_rle(Y).size
    y_ratio = 100 * y_compressed_size / y_original_size
    stats['Y_wo'] = {
        'original_size': y_original_size,
        'compressed_size': y_compressed_size,
        'compression_ratio': y_ratio
    }
    
    cr_original_size = Cr.size 
    cr_compressed_size = compress_rle(Cr).size
    cr_ratio = 100 * cr_compressed_size / cr_original_size
    stats['Cr_wo'] = {
        'original_size': cr_original_size,
        'compressed_size': cr_compressed_size,
        'compression_ratio': cr_ratio
    }

    cb_original_size = Cb.size 
    cb_compressed_size = compress_rle(Cb).size
    cb_ratio = 100 * cb_compressed_size / cb_original_size
    stats['Cb_wo'] = {
        'original_size': cb_original_size,
        'compressed_size': cb_compressed_size,
        'compression_ratio': cb_ratio
    }

    total_original = y_original_size + cr_original_size + cb_original_size
    total_compressed = y_compressed_size + cr_compressed_size + cb_compressed_size
    total_ratio = 100 * total_compressed / total_original
    stats['Total_wo'] = {
        'original_size': total_original,
        'compressed_size': total_compressed,
        'compression_ratio': total_ratio
    }



    return ver2(Y_c, Cb_c, Cr_c, Y.shape, Ratio, QY, QC, stats)


def decompressJPEG(JPEG):

    Y_dct = decompress_rle(JPEG.Y)
    Cb_dct = decompress_rle(JPEG.Cb)
    Cr_dct = decompress_rle(JPEG.Cr)


    if JPEG.ChromaRatio == "4:2:2":
        cr_shape = (JPEG.shape[0], JPEG.shape[1] // 2)
    else:
        cr_shape = JPEG.shape


    Y = DecompressLayer(Y_dct, JPEG.shape, JPEG.QY)
    Cr_sub = DecompressLayer(Cr_dct, cr_shape, JPEG.QC)
    Cb_sub = DecompressLayer(Cb_dct, cr_shape, JPEG.QC)

    Cr = chromaUpsample(Cr_sub, JPEG.ChromaRatio)
    Cb = chromaUpsample(Cb_sub, JPEG.ChromaRatio)

    YCrCb = np.dstack([Y, Cr, Cb]).clip(0, 255).astype(np.uint8)
    return YCrCbtoRGB(YCrCb)


QY= np.array([
        [16, 11, 10, 16, 24,  40,  51,  61],
        [12, 12, 14, 19, 26,  58,  60,  55],
        [14, 13, 16, 24, 40,  57,  69,  56],
        [14, 17, 22, 29, 51,  87,  80,  62],
        [18, 22, 37, 56, 68,  109, 103, 77],
        [24, 36, 55, 64, 81,  104, 113, 92],
        [49, 64, 78, 87, 103, 121, 120, 101],
        [72, 92, 95, 98, 112, 100, 103, 99],
        ])

QC= np.array([
        [17, 18, 24, 47, 99, 99, 99, 99],
        [18, 21, 26, 66, 99, 99, 99, 99],
        [24, 26, 56, 99, 99, 99, 99, 99],
        [47, 66, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        ])


def add_stats_table(document, stats, mode_description):
    table = document.add_table(rows=5, cols=4) 
    table.style = 'Table Grid'

    
    headers = ["Layer", "Original Size", "Compressed Size", "Compression Ratio"]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    
    
    for i, layer in enumerate(['Y', 'Cr', 'Cb', 'Total']):
        if layer in stats:
            table.cell(i+1, 0).text = layer
            table.cell(i+1, 1).text = str(stats[layer]['original_size'])
            table.cell(i+1, 2).text = str(stats[layer]['compressed_size'])
            table.cell(i+1, 3).text = f"{stats[layer]['compression_ratio']:.2f}%"
    
    document.add_paragraph(f"Compression Statistics - {mode_description}")
    
    
    wo_layers = [k for k in stats.keys() if k.endswith('_wo')]
    if wo_layers:
        document.add_paragraph("RLE without layer compression:")
        
        wo_rows = len(wo_layers) + 1
        wo_table = document.add_table(rows=wo_rows, cols=4)
        wo_table.style = 'Table Grid'
        
        for i, header in enumerate(headers):
            wo_table.cell(0, i).text = header
        
        for i, layer in enumerate(wo_layers):
            wo_table.cell(i+1, 0).text = layer.replace('_wo', '')
            wo_table.cell(i+1, 1).text = str(stats[layer]['original_size'])
            wo_table.cell(i+1, 2).text = str(stats[layer]['compressed_size'])
            wo_table.cell(i+1, 3).text = f"{stats[layer]['compression_ratio']:.2f}%"
    

def layersComp(przed_rgb, po_rgb, document=None):
    fig, axs = plt.subplots(4, 2, sharey=True)
    fig.set_size_inches(6, 8)

    axs[0, 0].imshow(przed_rgb)
    axs[0, 0].set_title("orig RGB")
    axs[0, 0].axis("off")

    axs[0, 1].imshow(po_rgb)
    axs[0, 1].set_title("po dekompresji RGB")
    axs[0, 1].axis("off")

    przed_ycc = cv2.cvtColor(przed_rgb, cv2.COLOR_RGB2YCrCb)
    po_ycc = cv2.cvtColor(po_rgb, cv2.COLOR_RGB2YCrCb)

    labels = ["Y", "Cr", "Cb"]
    for i in range(3):
        axs[i + 1, 0].imshow(przed_ycc[:, :, i], cmap="gray")
        axs[i + 1, 0].set_title(f"orig {labels[i]}")
        axs[i + 1, 0].axis("off")

        axs[i + 1, 1].imshow(po_ycc[:, :, i], cmap="gray")
        axs[i + 1, 1].set_title(f"po dekompresji {labels[i]}")
        axs[i + 1, 1].axis("off")

    plt.tight_layout()
    if document:
        plt.tight_layout(pad=1.5)
        memfile = BytesIO()
        plt.savefig(memfile)
        document.add_picture(memfile,width=Inches(6))
        memfile.close()
        plt.close()
    else:
        plt.show()

def process_image(fragment, document):
    document.add_heading("4:4:4, neutral Quantization", level=3)
    compressed = compressJPEG(fragment)
    decompressed = decompressJPEG(compressed)
    layersComp(fragment, decompressed, document)
    add_stats_table(document, compressed.stats, "4:4:4, neutral Quantization")
    
    document.add_heading("4:2:2, neutral Quantization", level=3)
    compressed = compressJPEG(fragment, "4:2:2")
    decompressed = decompressJPEG(compressed)
    layersComp(fragment, decompressed, document)
    add_stats_table(document, compressed.stats, "4:2:2, neutral Quantization")
    
    document.add_heading("4:4:4, Lumination 50%  Quantization", level=3)
    compressed = compressJPEG(fragment, "4:4:4", QY, QC)
    decompressed = decompressJPEG(compressed)
    layersComp(fragment, decompressed, document)
    add_stats_table(document, compressed.stats, "4:4:4, Lumination 50%  Quantization")
    
    document.add_heading("4:2:2, Lumination 50% Quantization", level=3)
    compressed = compressJPEG(fragment, "4:2:2", QY, QC)
    decompressed = decompressJPEG(compressed)
    layersComp(fragment, decompressed, document)
    add_stats_table(document, compressed.stats, "4:2:2, Lumination 50% Quantization")
        
    
        
    
    
    
def process_tests():
    data = [
    {"image_name": "BIG_0001.jpg", "x": 290, "y": 316, "width": 128, "height": 128},
    {"image_name": "BIG_0001.jpg", "x": 3946, "y": 609, "width": 128, "height": 128},
    {"image_name": "BIG_0001.jpg", "x": 838, "y": 2623, "width": 128, "height": 128},
    {"image_name": "BIG_0001.jpg", "x": 3764, "y": 2851, "width": 128, "height": 128},
    {"image_name": "BIG_0002.jpg", "x": 706, "y": 1629, "width": 128, "height": 128},
    {"image_name": "BIG_0002.jpg", "x": 2520, "y": 1538, "width": 128, "height": 128},
    {"image_name": "BIG_0002.jpg", "x": 3874, "y": 2957, "width": 128, "height": 128},
    {"image_name": "BIG_0002.jpg", "x": 4405, "y": 1940, "width": 128, "height": 128},
    {"image_name": "BIG_0003.jpg", "x": 765, "y": 1995, "width": 128, "height": 128},
    {"image_name": "BIG_0003.jpg", "x": 247, "y": 2509, "width": 128, "height": 128},
    {"image_name": "BIG_0003.jpg", "x": 1171, "y": 3602, "width": 128, "height": 128},
    {"image_name": "BIG_0003.jpg", "x": 2605, "y": 3364, "width": 128, "height": 128},
    {"image_name": "BIG_0004.png", "x": 79, "y": 312, "width": 128, "height": 128},
    {"image_name": "BIG_0004.png", "x": 163, "y": 622, "width": 128, "height": 128},
    {"image_name": "BIG_0004.png", "x": 861, "y": 60, "width": 128, "height": 128},
    {"image_name": "BIG_0004.png", "x": 1684, "y": 506, "width": 128, "height": 128},
    ]
    df = pd.DataFrame(data)

    image_dir = "IMG"
    
    document = Document()
    document.add_heading('Laboratorium 8', 0)
    



    for i, row in df.iterrows():
        image_path = os.path.join(image_dir, row["image_name"])
        img = cv2.imread(image_path)
        img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    
        if img is None:
            print(f"Error reading {image_path}")
            continue

        x, y, w, h = row["x"], row["y"], row["width"], row["height"]
        fragment = img[y:y+h, x:x+w]
        heading_text = (
        f"{row['image_name']} fragment: x={x}, y={y}, width={w}, height={h}"
        )
        document.add_heading(heading_text, level=2)
        process_image(fragment, document)
    
    
    document.save("tarakhkalo_lab8.docx")
        

if __name__ == "__main__":

    process_tests()