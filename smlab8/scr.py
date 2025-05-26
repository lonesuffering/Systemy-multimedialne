import numpy as np
import cv2
import matplotlib.pyplot as plt
import scipy.fftpack
from docx import Document
from docx.shared import Inches
import os

# === CONSTANTS ===
QY = np.array([
    [16, 11, 10, 16, 24, 40, 51, 61],
    [12, 12, 14, 19, 26, 58, 60, 55],
    [14, 13, 16, 24, 40, 57, 69, 56],
    [14, 17, 22, 29, 51, 87, 80, 62],
    [18, 22, 37, 56, 68, 109, 103, 77],
    [24, 36, 55, 64, 81, 104, 113, 92],
    [49, 64, 78, 87, 103, 121, 120, 101],
    [72, 92, 95, 98, 112, 100, 103, 99]
])

QC = np.array([
    [17, 18, 24, 47, 99, 99, 99, 99],
    [18, 21, 26, 66, 99, 99, 99, 99],
    [24, 26, 56, 99, 99, 99, 99, 99],
    [47, 66, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99]
])

QN = np.ones((8, 8))

# === RLE ===
def rle_encode(data):
    encoded = []
    prev = data[0]
    count = 1
    for i in range(1, len(data)):
        if data[i] == prev:
            count += 1
        else:
            encoded.append((prev, count))
            prev = data[i]
            count = 1
    encoded.append((prev, count))
    return encoded

def rle_decode(data):
    decoded = []
    for val, count in data:
        decoded.extend([val] * count)
    return np.array(decoded)

# === HELPER CLASSES ===
class ver2:
    def __init__(self, Y, Cb, Cr, OGShape, Ratio="4:4:4", QY=np.ones((8, 8)), QC=np.ones((8, 8)),
                 Y_range=None, Cb_range=None, Cr_range=None):
        self.shape = OGShape
        self.Y = Y
        self.Cb = Cb
        self.Cr = Cr
        self.ChromaRatio = Ratio
        self.QY = QY
        self.QC = QC
        self.Y_range = Y_range
        self.Cb_range = Cb_range
        self.Cr_range = Cr_range

# === HELPER FUNCTIONS ===
def dct2(a):
    return scipy.fftpack.dct(scipy.fftpack.dct(a.astype(float), axis=0, norm='ortho'), axis=1, norm='ortho')

def idct2(a):
    return scipy.fftpack.idct(scipy.fftpack.idct(a.astype(float), axis=0, norm='ortho'), axis=1, norm='ortho')

def zigzag(A):
    template = np.array([
        [0, 1, 5, 6, 14, 15, 27, 28],
        [2, 4, 7, 13, 16, 26, 29, 42],
        [3, 8, 12, 17, 25, 30, 41, 43],
        [9, 11, 18, 24, 31, 40, 44, 53],
        [10, 19, 23, 32, 39, 45, 52, 54],
        [20, 22, 33, 38, 46, 51, 55, 60],
        [21, 34, 37, 47, 50, 56, 59, 61],
        [35, 36, 48, 49, 57, 58, 62, 63]
    ])
    if len(A.shape) == 1:
        B = np.zeros((8, 8))
        for r in range(8):
            for c in range(8):
                B[r, c] = A[template[r, c]]
    else:
        B = np.zeros((64,))
        for r in range(8):
            for c in range(8):
                B[template[r, c]] = A[r, c]
    return B

def CompressBlock(block, Q):
    block = block - 128
    dct_block = dct2(block)
    q_block = np.round(dct_block / Q).astype(int)
    zz = zigzag(q_block)
    return zz

def DecompressBlock(vector, Q):
    q_block = zigzag(vector)
    dct_block = q_block * Q
    block = idct2(dct_block)
    block = block + 128
    return block

def CompressLayer(L, Q):
    S = []
    for w in range(0, L.shape[0], 8):
        for k in range(0, L.shape[1], 8):
            block = L[w:w + 8, k:k + 8]
            S.append(CompressBlock(block, Q))
    return np.concatenate(S)

def DecompressLayer(S, Q, shape):
    L = np.zeros(shape)
    m = shape[1] // 8
    for idx, i in enumerate(range(0, S.shape[0], 64)):
        vector = S[i:i + 64]
        k = (idx % m) * 8
        w = (idx // m) * 8
        L[w:w + 8, k:k + 8] = DecompressBlock(vector, Q)
    return L

def chroma_subsample(layer, ratio):
    if ratio == "4:2:2":
        return layer[:, ::2]
    else:
        return layer

def chroma_upsample(layer, shape, ratio):
    if ratio == "4:2:2":
        layer = np.repeat(layer, 2, axis=1)
        return layer[:, :shape[1]]
    else:
        return layer

def CompressJPEG(RGB, Ratio="4:4:4", QY=np.ones((8, 8)), QC=np.ones((8, 8))):
    YCrCb = cv2.cvtColor(RGB, cv2.COLOR_RGB2YCrCb).astype(int)
    Y = YCrCb[:, :, 0]
    Cr = YCrCb[:, :, 1]
    Cb = YCrCb[:, :, 2]

    Cr = chroma_subsample(Cr, Ratio)
    Cb = chroma_subsample(Cb, Ratio)

    Ymin, Ymax = Y.min(), Y.max()
    Cbmin, Cbmax = Cb.min(), Cb.max()
    Crmin, Crmax = Cr.min(), Cr.max()

    Yc = rle_encode(CompressLayer(Y, QY))
    Cbc = rle_encode(CompressLayer(Cb, QC))
    Crc = rle_encode(CompressLayer(Cr, QC))

    JPEG = ver2(Yc, Cbc, Crc, RGB.shape, Ratio, QY, QC,
                (Ymin, Ymax), (Cbmin, Cbmax), (Crmin, Crmax))
    return JPEG

def DecompressJPEG(JPEG):
    shape = JPEG.shape

    Y = DecompressLayer(rle_decode(JPEG.Y), JPEG.QY, (shape[0], shape[1]))
    Cb_shape = (shape[0], shape[1]) if JPEG.ChromaRatio == "4:4:4" else (shape[0], shape[1] // 2)
    Cr_shape = (shape[0], shape[1]) if JPEG.ChromaRatio == "4:4:4" else (shape[0], shape[1] // 2)

    Cb = DecompressLayer(rle_decode(JPEG.Cb), JPEG.QC, Cb_shape)
    Cr = DecompressLayer(rle_decode(JPEG.Cr), JPEG.QC, Cr_shape)

    Cb = chroma_upsample(Cb, (shape[0], shape[1]), JPEG.ChromaRatio)
    Cr = chroma_upsample(Cr, (shape[0], shape[1]), JPEG.ChromaRatio)

    Ymin, Ymax = JPEG.Y_range
    Cbmin, Cbmax = JPEG.Cb_range
    Crmin, Crmax = JPEG.Cr_range

    Y = np.clip(np.interp(Y, (Y.min(), Y.max()), (Ymin, Ymax)), 0, 255)
    Cb = np.clip(np.interp(Cb, (Cb.min(), Cb.max()), (Cbmin, Cbmax)), 0, 255)
    Cr = np.clip(np.interp(Cr, (Cr.min(), Cr.max()), (Crmin, Crmax)), 0, 255)

    YCrCb = np.dstack([Y, Cr, Cb]).astype(np.uint8)
    RGB = cv2.cvtColor(YCrCb, cv2.COLOR_YCrCb2RGB)
    return RGB


# === PLOTTING RESULTS ===
def plot_compression_results(crop, results, title_base, save_path_prefix):
    crop_YCrCb = cv2.cvtColor(crop, cv2.COLOR_RGB2YCrCb)
    for idx, (label, decompressed) in enumerate(results):
        fig, axs = plt.subplots(4, 2, figsize=(8, 10))
        fig.suptitle(f"{title_base} - {label}", fontsize=14)
        decompressed_YCrCb = cv2.cvtColor(decompressed, cv2.COLOR_RGB2YCrCb)

        axs[0,0].imshow(crop)
        axs[0,0].set_title('Original RGB')
        axs[0,1].imshow(decompressed)
        axs[0,1].set_title('Decompressed RGB')

        axs[1,0].imshow(crop_YCrCb[:,:,0], cmap='gray')
        axs[1,0].set_title('Original Y')
        axs[1,1].imshow(decompressed_YCrCb[:,:,0], cmap='gray')
        axs[1,1].set_title('Decompressed Y')

        axs[2,0].imshow(crop_YCrCb[:,:,2], cmap='gray')
        axs[2,0].set_title('Original Cb')
        axs[2,1].imshow(decompressed_YCrCb[:,:,2], cmap='gray')
        axs[2,1].set_title('Decompressed Cb')

        axs[3,0].imshow(crop_YCrCb[:,:,1], cmap='gray')
        axs[3,0].set_title('Original Cr')
        axs[3,1].imshow(decompressed_YCrCb[:,:,1], cmap='gray')
        axs[3,1].set_title('Decompressed Cr')

        for ax in axs.ravel():
            ax.axis('off')

        plt.tight_layout()
        plt.subplots_adjust(top=0.92)
        save_path = f"{save_path_prefix}_{label.replace(':','').replace(' ','')}.png"
        fig.savefig(save_path)
        plt.close(fig)

# === REPORT CREATION ===
def create_report(image_sets, filename="JPEG_Report.docx"):
    doc = Document()
    doc.add_heading('Raport Kompresji JPEG', 0)
    for img_idx, (img_name, crops) in enumerate(image_sets):
        doc.add_heading(f'Obraz {img_idx + 1}: {img_name}', level=1)
        for crop_idx, (crop, results) in enumerate(crops):
            doc.add_heading(f'Fragment {crop_idx + 1}', level=2)
            plot_base = f"plots/plot_{img_name}_fragment{crop_idx + 1}"
            plot_compression_results(crop, results, f'{img_name} - Fragment {crop_idx + 1}', plot_base)
            for label, _ in results:
                plot_path = f"{plot_base}_{label.replace(':','').replace(' ','')}.png"
                doc.add_picture(plot_path, width=Inches(5.5))
    doc.save(filename)

# === MAIN ===
if __name__ == "__main__":
    os.makedirs("fragments", exist_ok=True)
    os.makedirs("plots", exist_ok=True)
    images = []
    img = cv2.cvtColor(cv2.imread("BIG_0003.jpg"), cv2.COLOR_BGR2RGB)
    images.append(("BIG_0003", img))

    np.random.seed(42)
    image_sets = []
    for img_name, img in images:
        crops = []
        h, w = img.shape[:2]
        positions = [(np.random.randint(0, w - 128), np.random.randint(0, h - 128)) for _ in range(3)]
        for idx, (x, y) in enumerate(positions):
            crop = img[y:y + 128, x:x + 128, :]
            cv2.imwrite(f"fragments/fragment_{img_name}_{idx + 1}.png", cv2.cvtColor(crop, cv2.COLOR_RGB2BGR))
            results = [
                ("4:4:4 neutral", DecompressJPEG(CompressJPEG(crop, "4:4:4", QN, QN))),
                ("4:4:4 standard", DecompressJPEG(CompressJPEG(crop, "4:4:4", QY, QC))),
                ("4:2:2 neutral", DecompressJPEG(CompressJPEG(crop, "4:2:2", QN, QN))),
                ("4:2:2 standard", DecompressJPEG(CompressJPEG(crop, "4:2:2", QY, QC)))
            ]
            crops.append((crop, results))
        image_sets.append((img_name, crops))

    create_report(image_sets)
    print("✅ Raport JPEG_Report.docx został wygenerowany!")
