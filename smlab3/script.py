import os
import numpy as np
import cv2
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Путь к папкам с изображениями
BIG_IMAGES_DIR = "IMG_BIG"
SMALL_IMAGES_DIR = "SMALL"

# Метод ближайшего соседа
def nearest_neighbor_resize(image, scale_factor):
    height, width = image.shape[:2]
    new_height, new_width = int(height * scale_factor), int(width * scale_factor)
    resized_image = np.zeros((new_height, new_width, 3), dtype=np.uint8)
    
    for y in range(new_height):
        for x in range(new_width):
            src_x = int(x / scale_factor)
            src_y = int(y / scale_factor)
            resized_image[y, x] = image[src_y, src_x]
    
    return resized_image

# Билинейная интерполяция
def bilinear_interpolation_resize(image, scale_factor):
    height, width = image.shape[:2]
    new_height, new_width = int(height * scale_factor), int(width * scale_factor)
    resized_image = np.zeros((new_height, new_width, 3), dtype=np.uint8)
    
    for y in range(new_height):
        for x in range(new_width):
            src_x = x / scale_factor
            src_y = y / scale_factor
            
            x1, y1 = int(src_x), int(src_y)
            x2, y2 = min(x1 + 1, width - 1), min(y1 + 1, height - 1)
            
            dx, dy = src_x - x1, src_y - y1
            
            top = (1 - dx) * image[y1, x1] + dx * image[y1, x2]
            bottom = (1 - dx) * image[y2, x1] + dx * image[y2, x2]
            
            resized_image[y, x] = ((1 - dy) * top + dy * bottom).astype(np.uint8)
    
    return resized_image

def average_resize(image, scale_factor):
    height, width = image.shape[:2]
    new_height, new_width = int(height * scale_factor), int(width * scale_factor)
    resized_image = np.zeros((new_height, new_width, 3), dtype=np.uint8)
    
    block_size_x = width // new_width
    block_size_y = height // new_height
    
    for y in range(new_height):
        for x in range(new_width):
            start_x = x * block_size_x
            start_y = y * block_size_y
            end_x = min(start_x + block_size_x, width)  # Учитываем границы изображения
            end_y = min(start_y + block_size_y, height)
            
            block = image[start_y:end_y, start_x:end_x]
            
            if block.size == 0:  # Проверяем, что блок не пустой
                resized_image[y, x] = 0  # Заполняем черным цветом
            else:
                resized_image[y, x] = np.mean(block, axis=(0, 1)).astype(np.uint8)
    
    return resized_image

# Уменьшение изображения методом взвешенного среднего
def weighted_average_resize(image, scale_factor):
    weights = np.array([[1, 2, 1],
                        [2, 4, 2],
                        [1, 2, 1]]) / 16.0
    
    height, width = image.shape[:2]
    new_height, new_width = int(height * scale_factor), int(width * scale_factor)
    resized_image = np.zeros((new_height, new_width, 3), dtype=np.uint8)
    
    block_size_x = width // new_width
    block_size_y = height // new_height
    
    for y in range(new_height):
        for x in range(new_width):
            start_x = x * block_size_x
            start_y = y * block_size_y
            end_x = min(start_x + block_size_x, width)  # Учитываем границы изображения
            end_y = min(start_y + block_size_y, height)
            
            block = image[start_y:end_y, start_x:end_x]
            
            if block.size == 0:  # Проверяем, что блок не пустой
                resized_image[y, x] = 0  # Заполняем черным цветом
                continue
            
            for c in range(3):  # Для каждого канала (R, G, B)
                filtered_channel = cv2.filter2D(block[:, :, c], -1, weights)
                mean_value = np.mean(filtered_channel)
                resized_image[y, x, c] = mean_value.astype(np.uint8)
    
    return resized_image

def median_resize(image, scale_factor):
    height, width = image.shape[:2]
    new_height, new_width = int(height * scale_factor), int(width * scale_factor)
    resized_image = np.zeros((new_height, new_width, 3), dtype=np.uint8)
    
    block_size_x = width // new_width
    block_size_y = height // new_height
    
    for y in range(new_height):
        for x in range(new_width):
            start_x = x * block_size_x
            start_y = y * block_size_y
            end_x = min(start_x + block_size_x, width)  # Учитываем границы изображения
            end_y = min(start_y + block_size_y, height)
            
            block = image[start_y:end_y, start_x:end_x]
            
            if block.size == 0:  # Проверяем, что блок не пустой
                resized_image[y, x] = 0  # Заполняем черным цветом
            else:
                resized_image[y, x] = np.median(block, axis=(0, 1)).astype(np.uint8)
    
    return resized_image

# Визуализация результатов
def visualize_results(original, resized, title):
    plt.figure(figsize=(10, 5))
    plt.subplot(1, 2, 1)
    plt.title("Original")
    plt.imshow(cv2.cvtColor(original, cv2.COLOR_BGR2RGB))
    plt.axis('off')
    
    plt.subplot(1, 2, 2)
    plt.title(title)
    plt.imshow(cv2.cvtColor(resized, cv2.COLOR_BGR2RGB))
    plt.axis('off')
    plt.tight_layout()
    plt.show()

# Сохранение изображения в поток BytesIO
def save_image_to_bytes(fig):
    img_stream = BytesIO()
    fig.savefig(img_stream, format='png', bbox_inches='tight', dpi=300)  # Сохраняем изображение в буфер
    img_stream.seek(0)  # Перемещаем указатель в начало потока
    return img_stream

RESIZE_METHODS = {
    "Nearest Neighbor": nearest_neighbor_resize,
    "Bilinear Interpolation": bilinear_interpolation_resize,
    "Average Resize": average_resize,
    "Weighted Average Resize": weighted_average_resize,
    "Median Resize": median_resize
}
# Создание отчета в формате .docx
def generate_report():
    document = Document()
    document.add_heading('Отчет по изменению размера изображений', 0)
    
    # Тестирование больших изображений
    big_images = [f for f in os.listdir(BIG_IMAGES_DIR) if f.endswith(('.jpg', '.png'))]
    for big_image_name in big_images:
        big_image_path = os.path.join(BIG_IMAGES_DIR, big_image_name)
        big_image = cv2.imread(big_image_path)
        
        if big_image is None:  # Проверяем, что изображение загружено
            print(f"Ошибка: Не удалось загрузить изображение {big_image_name}")
            continue
        
        document.add_heading(f'Большое изображение: {big_image_name}', level=1)
        
        # Выбираем несколько фрагментов
        fragments = [
            (20, 20, 100, 100),
            (150, 150, 250, 250),
            (300, 300, 400, 400)
        ]
        
        for fragment_coords in fragments:
            x1, y1, x2, y2 = fragment_coords
            fragment = big_image[y1:y2, x1:x2]
            
            if fragment.size == 0:  # Проверяем, что фрагмент не пустой
                print(f"Ошибка: Пустой фрагмент для координат ({x1}, {y1}) - ({x2}, {y2})")
                continue
            
            document.add_heading(f'Фрагмент: ({x1}, {y1}) - ({x2}, {y2})', level=2)
            
            for method_name, method in RESIZE_METHODS.items():
                for scale in [0.5, 0.1]:
                    resized_fragment = method(fragment, scale)
                    
                    # Добавляем оригинальное изображение и результат в отчет
                    fig, ax = plt.subplots(1, 2, figsize=(10, 5))
                    ax[0].imshow(cv2.cvtColor(fragment, cv2.COLOR_BGR2RGB))
                    ax[0].set_title('Оригинал')
                    ax[0].axis('off')
                    
                    ax[1].imshow(cv2.cvtColor(resized_fragment, cv2.COLOR_BGR2RGB))
                    ax[1].set_title(f'{method_name} (Scale: {scale})')
                    ax[1].axis('off')
                    
                    plt.tight_layout()
                    
                    # Сохраняем изображение в поток BytesIO
                    img_stream = save_image_to_bytes(fig)
                    document.add_picture(img_stream, width=Inches(6))
                    
                    plt.close(fig)
    
    # Тестирование малых изображений
    small_images = [f for f in os.listdir(SMALL_IMAGES_DIR) if f.endswith(('.tif', '.png', '.jpg'))]
    for small_image_name in small_images:
        small_image_path = os.path.join(SMALL_IMAGES_DIR, small_image_name)
        small_image = cv2.imread(small_image_path)
        
        if small_image is None:  # Проверяем, что изображение загружено
            print(f"Ошибка: Не удалось загрузить изображение {small_image_name}")
            continue
        
        document.add_heading(f'Маленькое изображение: {small_image_name}', level=1)
        
        for method_name, method in RESIZE_METHODS.items():
            for scale in [3, 5]:
                resized_image = method(small_image, scale)
                
                # Добавляем оригинальное изображение и результат в отчет
                fig, ax = plt.subplots(1, 2, figsize=(10, 5))
                ax[0].imshow(cv2.cvtColor(small_image, cv2.COLOR_BGR2RGB))
                ax[0].set_title('Оригинал')
                ax[0].axis('off')
                
                ax[1].imshow(cv2.cvtColor(resized_image, cv2.COLOR_BGR2RGB))
                ax[1].set_title(f'{method_name} (Scale: {scale})')
                ax[1].axis('off')
                
                plt.tight_layout()
                
                # Сохраняем изображение в поток BytesIO
                img_stream = save_image_to_bytes(fig)
                document.add_picture(img_stream, width=Inches(6))
                
                plt.close(fig)
    
    # Сохраняем отчет
    document.save('report.docx')

# Основной блок выполнения
if __name__ == "__main__":
    generate_report()